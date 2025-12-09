#!/usr/bin/env python3
"""
Generate sample compliance documents for testing the transformation pipeline.
These are synthetic documents with realistic FedRAMP/NIST 800-53 content.
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os

OUTPUT_DIR = "docs/sample"

def create_access_control_policy():
    """Create a sample Access Control Policy document."""
    doc = Document()

    # Title
    title = doc.add_heading("Access Control Policy", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata
    doc.add_paragraph(f"Version: 1.0")
    doc.add_paragraph(f"Effective Date: {datetime.now().strftime('%Y-%m-%d')}")
    doc.add_paragraph("Classification: Internal Use Only")
    doc.add_paragraph("")

    # Purpose
    doc.add_heading("1. Purpose", level=1)
    doc.add_paragraph(
        "This Access Control Policy establishes the requirements for controlling access "
        "to information systems and data in accordance with NIST 800-53 controls and "
        "FedRAMP requirements."
    )

    # Scope
    doc.add_heading("2. Scope", level=1)
    doc.add_paragraph(
        "This policy applies to all employees, contractors, and third-party users who "
        "access organizational information systems."
    )

    # Policy Statements
    doc.add_heading("3. Policy Statements", level=1)

    # AC-1
    doc.add_heading("3.1 Access Control Policy and Procedures (AC-1)", level=2)
    doc.add_paragraph(
        "The organization develops, documents, and disseminates an access control policy "
        "that addresses purpose, scope, roles, responsibilities, and compliance. This policy "
        "is reviewed and updated annually or when significant changes occur."
    )

    # AC-2
    doc.add_heading("3.2 Account Management (AC-2)", level=2)
    doc.add_paragraph(
        "The organization manages information system accounts, including establishing, "
        "activating, modifying, disabling, and removing accounts. Account management includes:"
    )
    bullets = [
        "Identifying authorized users and access privileges (AC-2.a)",
        "Requiring manager approval for account creation (AC-2.b)",
        "Establishing conditions for group membership (AC-2.c)",
        "Specifying authorized users and access types (AC-2.d)",
        "Reviewing accounts every 90 days (AC-2.j)"
    ]
    for bullet in bullets:
        doc.add_paragraph(bullet, style='List Bullet')

    # AC-3
    doc.add_heading("3.3 Access Enforcement (AC-3)", level=2)
    doc.add_paragraph(
        "The information system enforces approved authorizations for logical access "
        "to information and system resources in accordance with applicable access "
        "control policies. Role-based access control (RBAC) is implemented across all systems."
    )

    # AC-4
    doc.add_heading("3.4 Information Flow Enforcement (AC-4)", level=2)
    doc.add_paragraph(
        "The information system enforces approved authorizations for controlling the "
        "flow of information within the system and between interconnected systems "
        "based on applicable policy."
    )

    # AC-5
    doc.add_heading("3.5 Separation of Duties (AC-5)", level=2)
    doc.add_paragraph(
        "The organization separates duties of individuals as necessary to prevent "
        "malevolent activity without collusion. Separation of duties includes:"
    )
    sep_duties = [
        "Dividing mission functions and distinct information system support functions",
        "Separating development, testing, and production environments",
        "Requiring multiple individuals for critical operations"
    ]
    for duty in sep_duties:
        doc.add_paragraph(duty, style='List Bullet')

    # AC-6
    doc.add_heading("3.6 Least Privilege (AC-6)", level=2)
    doc.add_paragraph(
        "The organization employs the principle of least privilege, allowing only "
        "authorized accesses for users and processes which are necessary to accomplish "
        "assigned tasks in accordance with organizational missions and business functions."
    )
    doc.add_paragraph(
        "Enhancement AC-6(1): The organization explicitly authorizes access to security "
        "functions and security-relevant information."
    )
    doc.add_paragraph(
        "Enhancement AC-6(2): The organization requires that users of information system "
        "accounts with access to security functions use non-privileged accounts when "
        "accessing other system functions."
    )

    # Roles and Responsibilities
    doc.add_heading("4. Roles and Responsibilities", level=1)

    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'

    headers = table.rows[0].cells
    headers[0].text = "Role"
    headers[1].text = "Responsibility"

    rows_data = [
        ("Chief Information Security Officer (CISO)", "Oversight of access control program"),
        ("System Administrators", "Implementation and monitoring of access controls"),
        ("All Users", "Compliance with access control policies")
    ]

    for i, (role, resp) in enumerate(rows_data, 1):
        row = table.rows[i].cells
        row[0].text = role
        row[1].text = resp

    # Compliance
    doc.add_heading("5. Compliance", level=1)
    doc.add_paragraph(
        "Non-compliance with this policy may result in disciplinary action up to and "
        "including termination. Violations should be reported to the CISO immediately."
    )

    # Save
    filepath = os.path.join(OUTPUT_DIR, "sample-access-control-policy.docx")
    doc.save(filepath)
    print(f"Created: {filepath}")


def create_ssp_excerpt():
    """Create a sample SSP excerpt with multiple control families."""
    doc = Document()

    title = doc.add_heading("System Security Plan - Control Implementation", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("System Name: CloudSecure Platform")
    doc.add_paragraph("FedRAMP Impact Level: Moderate")
    doc.add_paragraph(f"Document Version: 2.1")
    doc.add_paragraph(f"Last Updated: {datetime.now().strftime('%Y-%m-%d')}")
    doc.add_paragraph("")

    # Control Implementation Details
    doc.add_heading("Control Implementation Summary", level=1)

    controls = [
        {
            "id": "AC-2",
            "title": "Account Management",
            "status": "Implemented",
            "description": "Account management is implemented through Okta identity provider with automated provisioning and deprovisioning workflows. Quarterly access reviews are conducted using ServiceNow."
        },
        {
            "id": "AU-2",
            "title": "Audit Events",
            "status": "Implemented",
            "description": "The system audits login/logout events, privilege escalation, configuration changes, and data access. Logs are forwarded to Splunk SIEM for analysis and retention."
        },
        {
            "id": "AU-6",
            "title": "Audit Review, Analysis, and Reporting",
            "status": "Implemented",
            "description": "Security analysts review audit logs daily using Splunk dashboards. Automated alerts notify the SOC of suspicious activities. Weekly reports are provided to management."
        },
        {
            "id": "CM-2",
            "title": "Baseline Configuration",
            "status": "Implemented",
            "description": "Baseline configurations are maintained in Terraform and Ansible. All infrastructure changes go through GitOps workflows with automated compliance checks."
        },
        {
            "id": "CM-6",
            "title": "Configuration Settings",
            "status": "Implemented",
            "description": "CIS Benchmarks are applied to all systems. Configuration compliance is monitored continuously using AWS Config and custom Lambda functions."
        },
        {
            "id": "IA-2",
            "title": "Identification and Authentication",
            "status": "Implemented",
            "description": "Multi-factor authentication is required for all users. PIV/CAC cards are supported for federal users. Passwords meet NIST 800-63B requirements."
        },
        {
            "id": "IA-2(1)",
            "title": "Network Access to Privileged Accounts",
            "status": "Implemented",
            "description": "MFA is enforced for all privileged access including admin consoles, SSH, and API access. Hardware tokens are required for production access."
        },
        {
            "id": "SC-7",
            "title": "Boundary Protection",
            "status": "Implemented",
            "description": "AWS WAF and CloudFront provide boundary protection. VPC network ACLs and security groups enforce microsegmentation. All traffic is logged and analyzed."
        },
        {
            "id": "SC-8",
            "title": "Transmission Confidentiality and Integrity",
            "status": "Implemented",
            "description": "All data in transit is encrypted using TLS 1.3. Certificate management is automated through AWS Certificate Manager. HSTS is enforced on all endpoints."
        },
        {
            "id": "SC-13",
            "title": "Cryptographic Protection",
            "status": "Implemented",
            "description": "FIPS 140-2 validated modules are used for all cryptographic operations. AWS KMS with FIPS endpoints provides key management."
        },
        {
            "id": "SI-2",
            "title": "Flaw Remediation",
            "status": "Implemented",
            "description": "Vulnerability scanning is performed weekly using Qualys. Critical vulnerabilities are remediated within 15 days, high within 30 days per FedRAMP requirements."
        },
        {
            "id": "SI-4",
            "title": "Information System Monitoring",
            "status": "Implemented",
            "description": "CloudWatch, GuardDuty, and Splunk provide continuous monitoring. IDS/IPS capabilities detect and block malicious activity. 24/7 SOC coverage is provided."
        }
    ]

    # Create table
    table = doc.add_table(rows=len(controls)+1, cols=4)
    table.style = 'Table Grid'

    # Headers
    headers = table.rows[0].cells
    headers[0].text = "Control ID"
    headers[1].text = "Control Title"
    headers[2].text = "Status"
    headers[3].text = "Implementation Description"

    for i, ctrl in enumerate(controls, 1):
        row = table.rows[i].cells
        row[0].text = ctrl["id"]
        row[1].text = ctrl["title"]
        row[2].text = ctrl["status"]
        row[3].text = ctrl["description"]

    # Save
    filepath = os.path.join(OUTPUT_DIR, "sample-ssp-excerpt.docx")
    doc.save(filepath)
    print(f"Created: {filepath}")


def create_poam():
    """Create a sample Plan of Action and Milestones document."""
    doc = Document()

    title = doc.add_heading("Plan of Action and Milestones (POA&M)", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("System Name: CloudSecure Platform")
    doc.add_paragraph(f"Report Date: {datetime.now().strftime('%Y-%m-%d')}")
    doc.add_paragraph("Prepared By: Security Operations Team")
    doc.add_paragraph("")

    doc.add_heading("Open Findings", level=1)

    findings = [
        {
            "id": "POAM-2024-001",
            "control": "AC-2",
            "title": "Incomplete Quarterly Access Reviews",
            "severity": "Moderate",
            "status": "In Progress",
            "due": "2025-02-15",
            "owner": "IAM Team Lead",
            "description": "Q3 access reviews did not cover contractor accounts.",
            "remediation": "Expand review scope to include all account types. Automate review process."
        },
        {
            "id": "POAM-2024-002",
            "control": "CM-6",
            "title": "CIS Benchmark Deviations",
            "severity": "Low",
            "status": "Open",
            "due": "2025-03-01",
            "owner": "Platform Engineering",
            "description": "15 servers have documented deviations from CIS benchmarks.",
            "remediation": "Review deviations for business justification. Remediate or document exceptions."
        },
        {
            "id": "POAM-2024-003",
            "control": "RA-5",
            "title": "Delayed Critical Vulnerability Remediation",
            "severity": "High",
            "status": "In Progress",
            "due": "2025-01-31",
            "owner": "Security Engineering",
            "description": "CVE-2024-12345 remediation exceeded 15-day SLA.",
            "remediation": "Implement automated patching for critical vulnerabilities."
        },
        {
            "id": "POAM-2024-004",
            "control": "AU-6",
            "title": "Incomplete Log Review Documentation",
            "severity": "Low",
            "status": "Open",
            "due": "2025-02-28",
            "owner": "SOC Manager",
            "description": "Weekly log review reports missing for November 2024.",
            "remediation": "Automate report generation. Implement review checklist."
        },
        {
            "id": "POAM-2024-005",
            "control": "IR-4",
            "title": "Incident Response Procedure Updates",
            "severity": "Moderate",
            "status": "Open",
            "due": "2025-02-01",
            "owner": "CISO",
            "description": "IR procedures do not reflect new cloud infrastructure.",
            "remediation": "Update IR playbooks for AWS environment. Conduct tabletop exercise."
        }
    ]

    # Create table
    table = doc.add_table(rows=len(findings)+1, cols=6)
    table.style = 'Table Grid'

    headers = table.rows[0].cells
    headers[0].text = "POA&M ID"
    headers[1].text = "Control"
    headers[2].text = "Finding"
    headers[3].text = "Severity"
    headers[4].text = "Status"
    headers[5].text = "Due Date"

    for i, finding in enumerate(findings, 1):
        row = table.rows[i].cells
        row[0].text = finding["id"]
        row[1].text = finding["control"]
        row[2].text = finding["title"]
        row[3].text = finding["severity"]
        row[4].text = finding["status"]
        row[5].text = finding["due"]

    doc.add_paragraph("")

    # Detailed findings
    doc.add_heading("Finding Details", level=1)

    for finding in findings:
        doc.add_heading(f"{finding['id']}: {finding['title']}", level=2)
        doc.add_paragraph(f"Control: {finding['control']}")
        doc.add_paragraph(f"Severity: {finding['severity']}")
        doc.add_paragraph(f"Status: {finding['status']}")
        doc.add_paragraph(f"Due Date: {finding['due']}")
        doc.add_paragraph(f"Owner: {finding['owner']}")
        doc.add_paragraph(f"Description: {finding['description']}")
        doc.add_paragraph(f"Remediation Plan: {finding['remediation']}")
        doc.add_paragraph("")

    # Save
    filepath = os.path.join(OUTPUT_DIR, "sample-poam.docx")
    doc.save(filepath)
    print(f"Created: {filepath}")


def create_ir_procedure():
    """Create a sample Incident Response procedure document."""
    doc = Document()

    title = doc.add_heading("Incident Response Procedure", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("Document ID: IR-PROC-001")
    doc.add_paragraph("Version: 3.2")
    doc.add_paragraph(f"Effective Date: {datetime.now().strftime('%Y-%m-%d')}")
    doc.add_paragraph("Classification: Internal Use Only")
    doc.add_paragraph("")

    # Purpose
    doc.add_heading("1. Purpose", level=1)
    doc.add_paragraph(
        "This procedure establishes the incident response process in accordance with "
        "NIST 800-53 IR controls and FedRAMP requirements. It ensures consistent "
        "handling of security incidents to minimize impact and enable recovery."
    )

    # Scope
    doc.add_heading("2. Scope", level=1)
    doc.add_paragraph(
        "This procedure applies to all security incidents affecting the CloudSecure "
        "Platform and its components, including infrastructure, applications, and data."
    )

    # Applicable Controls
    doc.add_heading("3. Applicable Controls", level=1)
    controls = [
        "IR-1: Incident Response Policy and Procedures",
        "IR-2: Incident Response Training",
        "IR-3: Incident Response Testing",
        "IR-4: Incident Handling",
        "IR-5: Incident Monitoring",
        "IR-6: Incident Reporting",
        "IR-7: Incident Response Assistance",
        "IR-8: Incident Response Plan"
    ]
    for ctrl in controls:
        doc.add_paragraph(ctrl, style='List Bullet')

    # Incident Categories
    doc.add_heading("4. Incident Categories (IR-4)", level=1)

    categories = [
        ("Category 1 - Critical", "Active exploitation, data breach, system compromise"),
        ("Category 2 - High", "Attempted intrusion, malware detection, policy violation"),
        ("Category 3 - Medium", "Suspicious activity, vulnerability discovery"),
        ("Category 4 - Low", "Security queries, false positives, informational")
    ]

    table = doc.add_table(rows=len(categories)+1, cols=2)
    table.style = 'Table Grid'
    headers = table.rows[0].cells
    headers[0].text = "Category"
    headers[1].text = "Description"

    for i, (cat, desc) in enumerate(categories, 1):
        row = table.rows[i].cells
        row[0].text = cat
        row[1].text = desc

    doc.add_paragraph("")

    # Response Procedure
    doc.add_heading("5. Incident Response Procedure", level=1)

    doc.add_heading("5.1 Detection and Analysis (IR-4.a)", level=2)
    steps = [
        "Monitor security alerts from Splunk, GuardDuty, and other sources",
        "Validate alert and determine if incident is real or false positive",
        "Categorize incident based on severity and type",
        "Document initial findings in incident tracking system",
        "Notify appropriate personnel based on category"
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {step}")

    doc.add_heading("5.2 Containment (IR-4.b)", level=2)
    steps = [
        "Isolate affected systems if necessary",
        "Preserve evidence for forensic analysis",
        "Implement temporary mitigations",
        "Block malicious IPs or domains",
        "Revoke compromised credentials"
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {step}")

    doc.add_heading("5.3 Eradication and Recovery (IR-4.c)", level=2)
    steps = [
        "Remove malware or unauthorized access",
        "Apply patches or configuration changes",
        "Restore systems from clean backups if needed",
        "Verify system integrity before returning to production",
        "Monitor for recurrence"
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {step}")

    doc.add_heading("5.4 Post-Incident Activity (IR-4.d)", level=2)
    steps = [
        "Conduct lessons learned meeting within 5 business days",
        "Document incident timeline and response actions",
        "Identify process improvements",
        "Update playbooks and procedures as needed",
        "Report to FedRAMP PMO within required timeframes (IR-6)"
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {step}")

    # Reporting Requirements
    doc.add_heading("6. Reporting Requirements (IR-6)", level=1)

    reporting = [
        ("US-CERT", "Within 1 hour for Cat 1, 72 hours for Cat 2"),
        ("FedRAMP PMO", "Within 24 hours for significant incidents"),
        ("Agency AO", "Within 24 hours for incidents affecting agency data"),
        ("Internal CISO", "Immediately for all categories")
    ]

    table = doc.add_table(rows=len(reporting)+1, cols=2)
    table.style = 'Table Grid'
    headers = table.rows[0].cells
    headers[0].text = "Reporting Entity"
    headers[1].text = "Timeline"

    for i, (entity, timeline) in enumerate(reporting, 1):
        row = table.rows[i].cells
        row[0].text = entity
        row[1].text = timeline

    # Contact Information
    doc.add_heading("7. Contact Information", level=1)
    doc.add_paragraph("Security Operations Center (SOC): soc@example.com")
    doc.add_paragraph("CISO: ciso@example.com")
    doc.add_paragraph("Incident Response Team: irt@example.com")
    doc.add_paragraph("24/7 Hotline: 1-800-SEC-INCI")

    # Save
    filepath = os.path.join(OUTPUT_DIR, "sample-ir-procedure.docx")
    doc.save(filepath)
    print(f"Created: {filepath}")


if __name__ == "__main__":
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    print("Creating sample compliance documents...")
    create_access_control_policy()
    create_ssp_excerpt()
    create_poam()
    create_ir_procedure()
    print("Done!")
